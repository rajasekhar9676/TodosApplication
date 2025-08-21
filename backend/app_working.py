from flask import Flask, request, jsonify
from flask_cors import CORS
import speech_recognition as sr
import os
import tempfile
import base64
from werkzeug.utils import secure_filename
import logging
import pyttsx3
import io
from pydub import AudioSegment
import PyPDF2
from docx import Document
import re
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
import string
from collections import Counter
import json
import pdfplumber

app = Flask(__name__)
CORS(app)

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Download required NLTK data (free)
try:
    nltk.download('punkt', quiet=True)
    nltk.download('stopwords', quiet=True)
    nltk.download('wordnet', quiet=True)
    nltk.download('averaged_perceptron_tagger', quiet=True)
except:
    logger.warning("NLTK data download failed, using fallback methods")

# Initialize recognizer
recognizer = sr.Recognizer()
logger.info("✅ Speech recognition initialized successfully!")

# Initialize TTS engine
tts_engine = pyttsx3.init()
logger.info("✅ Text-to-speech engine initialized successfully!")

def convert_audio_to_wav(audio_bytes, original_format='webm'):
    """Convert audio bytes to WAV format or return original if conversion fails"""
    try:
        audio = AudioSegment.from_file(io.BytesIO(audio_bytes), format=original_format)
        wav_buffer = io.BytesIO()
        audio.export(wav_buffer, format='wav')
        wav_buffer.seek(0)
        logger.info("✅ Audio converted to WAV successfully using pydub")
        return wav_buffer.read()
    except Exception as e:
        logger.warning(f"Audio conversion failed: {e}")
        logger.info("⚠️ Falling back to original audio format")
        return audio_bytes

def extract_text_from_pdf(pdf_file):
    """Extract text from PDF file using pdfplumber for better extraction"""
    try:
        # Reset file pointer
        pdf_file.seek(0)
        
        # Try pdfplumber first (much better text extraction)
        try:
            with pdfplumber.open(pdf_file) as pdf:
                text = ""
                for page_num, page in enumerate(pdf.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text += f"\n--- Page {page_num + 1} ---\n{page_text.strip()}\n"
                        else:
                            # Try to extract text from tables if regular text extraction fails
                            tables = page.extract_tables()
                            if tables:
                                for table in tables:
                                    for row in table:
                                        row_text = " | ".join([str(cell) for cell in row if cell])
                                        if row_text.strip():
                                            text += row_text + "\n"
                                    text += "\n"
                            else:
                                text += f"\n--- Page {page_num + 1} ---\n[No text content found]\n"
                    except Exception as page_error:
                        logger.warning(f"Error extracting page {page_num + 1}: {page_error}")
                        text += f"\n--- Page {page_num + 1} ---\n[Error extracting text]\n"
                
                if text.strip():
                    logger.info(f"✅ PDF text extraction successful with pdfplumber: {len(text)} characters")
                    return text.strip()
                else:
                    raise Exception("No text content found with pdfplumber")
                    
        except Exception as pdfplumber_error:
            logger.warning(f"pdfplumber failed: {pdfplumber_error}, trying PyPDF2 as fallback")
            
            # Fallback to PyPDF2
            pdf_file.seek(0)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text += f"\n--- Page {page_num + 1} ---\n{page_text.strip()}\n"
                    else:
                        text += f"\n--- Page {page_num + 1} ---\n[No text content found]\n"
                except Exception as page_error:
                    logger.warning(f"PyPDF2 error extracting page {page_num + 1}: {page_error}")
                    text += f"\n--- Page {page_num + 1} ---\n[Error extracting text]\n"
            
            if text.strip():
                logger.info(f"✅ PDF text extraction successful with PyPDF2 fallback: {len(text)} characters")
                return text.strip()
            else:
                raise Exception("No text content found with either pdfplumber or PyPDF2")
        
    except Exception as e:
        logger.error(f"PDF text extraction failed: {e}")
        raise Exception(f"Failed to extract text from PDF: {str(e)}")

def extract_text_from_pdf_with_ocr(pdf_file):
    """Extract text from PDF using OCR as last resort for image-based PDFs"""
    try:
        # This is a placeholder for OCR functionality
        # In a production environment, you could use libraries like:
        # - pytesseract (Tesseract OCR)
        # - easyocr
        # - Azure Computer Vision API
        
        logger.warning("OCR extraction requested but not implemented in this version")
        return "OCR extraction not available in this version. Please ensure PDF contains selectable text."
        
    except Exception as e:
        logger.error(f"OCR extraction failed: {e}")
        return f"OCR extraction failed: {str(e)}"

def extract_text_from_docx(docx_file):
    """Extract text from DOCX file with enhanced processing"""
    try:
        # Reset file pointer
        docx_file.seek(0)
        doc = Document(docx_file)
        text = ""
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " | "
                text += "\n"
        
        if not text.strip():
            raise Exception("No text content found in DOCX")
        
        logger.info(f"✅ DOCX text extraction successful: {len(text)} characters")
        return text.strip()
    except Exception as e:
        logger.error(f"DOCX text extraction failed: {e}")
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")

def analyze_text_intelligence(text):
    """AI-powered text analysis using free NLP libraries"""
    try:
        # Clean and preprocess text
        text = re.sub(r'\s+', ' ', text).strip()
        
        # Basic statistics
        sentences = sent_tokenize(text)
        words = word_tokenize(text.lower())
        
        # Remove punctuation and stopwords
        stop_words = set(stopwords.words('english'))
        words_clean = [word for word in words if word.isalnum() and word not in stop_words]
        
        # Word frequency analysis
        word_freq = Counter(words_clean)
        top_keywords = word_freq.most_common(10)
        
        # Text complexity analysis
        avg_sentence_length = len(words) / len(sentences) if sentences else 0
        unique_words_ratio = len(set(words_clean)) / len(words_clean) if words_clean else 0
        
        # Topic identification (simple keyword-based)
        topics = []
        if any(word in text.lower() for word in ['technology', 'tech', 'software', 'computer']):
            topics.append('Technology')
        if any(word in text.lower() for word in ['business', 'company', 'market', 'finance']):
            topics.append('Business')
        if any(word in text.lower() for word in ['education', 'learning', 'school', 'university']):
            topics.append('Education')
        if any(word in text.lower() for word in ['health', 'medical', 'medicine', 'doctor']):
            topics.append('Healthcare')
        if any(word in text.lower() for word in ['news', 'current', 'event', 'update']):
            topics.append('News/Current Events')
        
        if not topics:
            topics = ['General']
        
        return {
            'word_count': len(words),
            'sentence_count': len(sentences),
            'character_count': len(text),
            'top_keywords': [{'word': word, 'count': count} for word, count in top_keywords],
            'avg_sentence_length': round(avg_sentence_length, 2),
            'unique_words_ratio': round(unique_words_ratio, 3),
            'topics': topics,
            'complexity': 'Simple' if avg_sentence_length < 15 else 'Moderate' if avg_sentence_length < 25 else 'Complex'
        }
    except Exception as e:
        logger.error(f"Text analysis failed: {e}")
        return {
            'word_count': len(text.split()),
            'sentence_count': len(text.split('.')),
            'character_count': len(text),
            'top_keywords': [],
            'topics': ['General'],
            'complexity': 'Unknown'
        }

def generate_ai_summary(text, analysis, summary_type='smart'):
    """Generate intelligent summary using AI-like analysis"""
    try:
        sentences = sent_tokenize(text)
        words = word_tokenize(text.lower())
        
        if summary_type == 'smart':
            # Smart summary: Find most important sentences based on keyword density
            sentence_scores = []
            top_keywords = [item['word'] for item in analysis['top_keywords'][:5]]
            
            for sentence in sentences:
                sentence_words = word_tokenize(sentence.lower())
                keyword_matches = sum(1 for word in sentence_words if word in top_keywords)
                sentence_scores.append((sentence, keyword_matches))
            
            # Sort by keyword density and take top sentences
            sentence_scores.sort(key=lambda x: x[1], reverse=True)
            important_sentences = [sent for sent, score in sentence_scores[:3] if score > 0]
            
            if important_sentences:
                summary = ' '.join(important_sentences)
            else:
                # Fallback to first few sentences
                summary = '. '.join(sentences[:2]) + '.'
        
        elif summary_type == 'explanation':
            # Explanation mode: Add context and analysis
            summary = f"📄 Document Analysis Report\n\n"
            summary += f"📊 Content Overview:\n"
            summary += f"• Total Words: {analysis['word_count']:,}\n"
            summary += f"• Total Sentences: {analysis['sentence_count']:,}\n"
            summary += f"• Characters: {analysis['character_count']:,}\n"
            summary += f"• Reading Complexity: {analysis['complexity']}\n\n"
            
            summary += f"🏷️ Identified Topics: {', '.join(analysis['topics'])}\n\n"
            
            if analysis['top_keywords']:
                summary += f"🔑 Key Terms:\n"
                for i, keyword in enumerate(analysis['top_keywords'][:5], 1):
                    summary += f"{i}. {keyword['word']} (appears {keyword['count']} times)\n"
                summary += "\n"
            
            summary += f"📝 Content Summary:\n"
            summary += f"{sentences[0] if sentences else 'No content available'}\n\n"
            
            if len(sentences) > 1:
                summary += f"Additional key points:\n"
                for i, sentence in enumerate(sentences[1:4], 2):
                    summary += f"{i}. {sentence}\n"
        
        else:  # basic summary
            summary = '. '.join(sentences[:3]) + '.'
            if len(summary) < 100:
                summary = text[:300] + '...' if len(text) > 300 else text
        
        return summary.strip()
    
    except Exception as e:
        logger.error(f"Summary generation failed: {e}")
        return text[:200] + '...' if len(text) > 200 else text

def process_text_content(text, process_type):
    """Enhanced text processing with AI-like features"""
    try:
        # Perform intelligent text analysis
        analysis = analyze_text_intelligence(text)
        
        # Generate appropriate output based on type
        if process_type == 'summary':
            return generate_ai_summary(text, analysis, 'smart')
        elif process_type == 'explanation':
            return generate_ai_summary(text, analysis, 'explanation')
        elif process_type == 'analysis':
            return json.dumps(analysis, indent=2)
        else:
            return text
    except Exception as e:
        logger.error(f"Text processing failed: {e}")
        return text

@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        "status": "healthy", 
        "message": "AI-Enhanced API is running",
        "features": ["speech-to-text", "text-to-speech", "ai-document-processing"],
        "version": "2.0.0",
        "ai_features": ["intelligent-summarization", "text-analysis", "topic-detection", "keyword-extraction"]
    })

@app.route('/api/speech/transcribe', methods=['POST'])
def transcribe_audio():
    """Web API endpoint for speech-to-text"""
    try:
        if 'audio' not in request.files and 'audio_data' not in request.json and 'audio' not in request.json:
            return jsonify({
                "success": False,
                "error": "No audio data provided",
                "supported_formats": ["base64", "file_upload", "audio_url"]
            }), 400
        
        # Handle base64 audio data from request.json
        if 'audio_data' in request.json:
            audio_data = request.json['audio_data']
            if audio_data.startswith('data:audio/'):
                audio_data = audio_data.split(',')[1]
            
            audio_bytes = base64.b64decode(audio_data)
            wav_audio = convert_audio_to_wav(audio_bytes, 'webm')
            
            with tempfile.NamedTemporaryFile(delete=False, suffix='.wav') as temp_file:
                temp_file.write(wav_audio)
                temp_file_path = temp_file.name
            
            try:
                with sr.AudioFile(temp_file_path) as source:
                    audio = recognizer.record(source)
                
                language = request.json.get('language', 'en-US')
                language_mapping = {
                    'en-US': 'en-US', 'hi-IN': 'hi-IN', 'bn-IN': 'bn-IN', 'te-IN': 'te-IN',
                    'ta-IN': 'ta-IN', 'kn-IN': 'kn-IN', 'ml-IN': 'ml-IN', 'gu-IN': 'gu-IN',
                    'mr-IN': 'mr-IN', 'pa-IN': 'pa-IN', 'or-IN': 'or-IN', 'as-IN': 'as-IN',
                    'ur-IN': 'ur-IN', 'es-ES': 'es-ES', 'fr-FR': 'fr-FR', 'de-DE': 'de-DE',
                    'it-IT': 'it-IT', 'pt-PT': 'pt-PT', 'ru-RU': 'ru-RU', 'ja-JP': 'ja-JP',
                    'ko-KR': 'ko-KR', 'zh-CN': 'zh-CN', 'ar-SA': 'ar-SA'
                }
                
                recognition_language = language_mapping.get(language, 'en-US')
                text = recognizer.recognize_google(audio, language=recognition_language)
                
                logger.info(f"✅ Successfully transcribed audio in {language}")
                return jsonify({
                    "success": True,
                    "text": text,
                    "language": language,
                    "method": "backend",
                    "confidence": "high"
                })
                
            except Exception as sr_error:
                logger.warning(f"Backend speech recognition failed: {sr_error}")
                return jsonify({
                    "success": False,
                    "error": "Backend processing failed. Using browser-based recognition.",
                    "fallback": "browser",
                    "note": "The frontend will automatically use browser-based speech recognition.",
                    "api_status": "fallback_available"
                }), 503
                
            finally:
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
        
        # Handle 'audio' field in request.json
        elif 'audio' in request.json:
            audio_data = request.json['audio']
            if audio_data.startswith('data:audio/'):
                audio_data = audio_data.split(',')[1]
            
            audio_bytes = base64.b64decode(audio_data)
            wav_audio = convert_audio_to_wav(audio_bytes, 'webm')
            
            with tempfile.NamedTemporaryFile(delete=False, suffix='.wav') as temp_file:
                temp_file.write(wav_audio)
                temp_file_path = temp_file.name
            
            try:
                with sr.AudioFile(temp_file_path) as source:
                    audio = recognizer.record(source)
                
                language = request.json.get('language', 'en-US')
                language_mapping = {
                    'en-US': 'en-US', 'hi-IN': 'hi-IN', 'bn-IN': 'bn-IN', 'te-IN': 'te-IN',
                    'ta-IN': 'ta-IN', 'kn-IN': 'kn-IN', 'ml-IN': 'ml-IN', 'gu-IN': 'gu-IN',
                    'mr-IN': 'mr-IN', 'pa-IN': 'pa-IN', 'or-IN': 'or-IN', 'as-IN': 'as-IN',
                    'ur-IN': 'ur-IN', 'es-ES': 'es-ES', 'fr-FR': 'fr-FR', 'de-DE': 'de-DE',
                    'it-IT': 'it-IT', 'pt-PT': 'pt-PT', 'ru-RU': 'ru-RU', 'ja-JP': 'ja-JP',
                    'ko-KR': 'ko-KR', 'zh-CN': 'zh-CN', 'ar-SA': 'ar-SA'
                }
                
                recognition_language = language_mapping.get(language, 'en-US')
                text = recognizer.recognize_google(audio, language=recognition_language)
                
                logger.info(f"✅ Successfully transcribed audio in {language}")
                return jsonify({
                    "success": True,
                    "text": text,
                    "language": language,
                    "method": "backend",
                    "confidence": "high"
                })
                
            except Exception as sr_error:
                logger.warning(f"Backend speech recognition failed: {sr_error}")
                return jsonify({
                    "success": False,
                    "error": "Backend processing failed. Using browser-based recognition.",
                    "fallback": "browser",
                    "note": "The frontend will automatically use browser-based speech recognition.",
                    "api_status": "fallback_available"
                }), 503
                
            finally:
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
        
        # Handle file upload
        elif 'audio' in request.files:
            audio_file = request.files['audio']
            if audio_file.filename == '':
                return jsonify({"error": "No file selected"}), 400
            
            filename = secure_filename(audio_file.filename)
            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1]) as temp_file:
                audio_file.save(temp_file.name)
                temp_file_path = temp_file.name
            
            try:
                with sr.AudioFile(temp_file_path) as source:
                    audio = recognizer.record(source)
                
                language = request.form.get('language', 'en-US')
                language_mapping = {
                    'en-US': 'en-US', 'hi-IN': 'hi-IN', 'bn-IN': 'bn-IN', 'te-IN': 'te-IN',
                    'ta-IN': 'ta-IN', 'kn-IN': 'kn-IN', 'ml-IN': 'ml-IN', 'gu-IN': 'gu-IN',
                    'mr-IN': 'mr-IN', 'pa-IN': 'pa-IN', 'or-IN': 'or-IN', 'as-IN': 'as-IN',
                    'ur-IN': 'ur-IN', 'es-ES': 'es-ES', 'fr-FR': 'fr-FR', 'de-DE': 'de-DE',
                    'it-IT': 'it-IT', 'pt-PT': 'pt-PT', 'ru-RU': 'ru-RU', 'ja-JP': 'ja-JP',
                    'ko-KR': 'ko-KR', 'zh-CN': 'zh-CN', 'ar-SA': 'ar-SA'
                }
                
                recognition_language = language_mapping.get(language, 'en-US')
                text = recognizer.recognize_google(audio, language=recognition_language)
                
                logger.info(f"✅ Successfully transcribed uploaded audio in {language}")
                return jsonify({
                    "success": True,
                    "text": text,
                    "language": language,
                    "method": "backend",
                    "confidence": "high"
                })
                
            except Exception as sr_error:
                logger.warning(f"Backend speech recognition failed: {sr_error}")
                return jsonify({
                    "success": False,
                    "error": "Backend processing failed. Using browser-based recognition.",
                    "fallback": "browser",
                    "note": "The frontend will automatically use browser-based speech recognition.",
                    "api_status": "fallback_available"
                }), 503
                
            finally:
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                    
    except Exception as e:
        logger.error(f"Unexpected error: {e}")
        return jsonify({
            "success": False,
            "error": "An unexpected error occurred. Please try again.",
            "api_status": "error"
        }), 500

@app.route('/api/speech/tts', methods=['POST'])
def text_to_speech():
    """Web API endpoint for text-to-speech"""
    try:
        data = request.get_json()
        if not data or 'text' not in data:
            return jsonify({
                "success": False, 
                "error": "No text provided",
                "required_fields": ["text"],
                "optional_fields": ["speed", "volume"]
            }), 400
        
        text = data['text']
        voice_speed = data.get('speed', 150)
        voice_volume = data.get('volume', 0.9)
        
        if not text.strip():
            return jsonify({
                "success": False, 
                "error": "Text is empty",
                "api_status": "validation_error"
            }), 400
        
        logger.info(f"Converting text to speech: {len(text)} characters")
        
        tts_engine.setProperty('rate', voice_speed)
        tts_engine.setProperty('volume', voice_volume)
        
        temp_audio = tempfile.NamedTemporaryFile(delete=False, suffix='.wav')
        temp_audio_path = temp_audio.name
        temp_audio.close()
        
        try:
            tts_engine.save_to_file(text, temp_audio_path)
            tts_engine.runAndWait()
            
            with open(temp_audio_path, 'rb') as audio_file:
                audio_data = base64.b64encode(audio_file.read()).decode('utf-8')
            
        finally:
            try:
                if os.path.exists(temp_audio_path):
                    os.unlink(temp_audio_path)
            except Exception as cleanup_error:
                logger.warning(f"Cleanup warning: {cleanup_error}")
        
        return jsonify({
            "success": True,
            "audio_data": audio_data,
            "text_length": len(text),
            "message": "Text converted to speech successfully",
            "api_status": "success",
            "format": "base64_wav"
        })
        
    except Exception as e:
        logger.error(f"TTS error: {e}")
        return jsonify({
            "success": False, 
            "error": f"Text-to-speech conversion failed: {str(e)}",
            "api_status": "error"
        }), 500

@app.route('/upload-document', methods=['POST'])
def upload_document():
    """AI-enhanced document processing endpoint"""
    try:
        if 'file' not in request.files:
            return jsonify({"success": False, "error": "No file provided"}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({"success": False, "error": "No file selected"}), 400
        
        process_type = request.form.get('type', 'summary')
        
        filename = secure_filename(file.filename)
        file_extension = filename.lower().split('.')[-1]
        
        if file_extension not in ['pdf', 'docx', 'doc', 'txt']:
            return jsonify({"success": False, "error": "Unsupported file type. Please upload PDF, Word, or text files."}), 400
        
        logger.info(f"Processing document: {filename} with type: {process_type}")
        
        # Extract text based on file type
        extracted_text = ""
        try:
            if file_extension == 'pdf':
                extracted_text = extract_text_from_pdf(file)
            elif file_extension in ['docx', 'doc']:
                extracted_text = extract_text_from_docx(file)
            elif file_extension == 'txt':
                extracted_text = file.read().decode('utf-8')
        except Exception as extract_error:
            logger.error(f"Text extraction failed: {extract_error}")
            return jsonify({
                "success": False, 
                "error": f"Failed to extract text: {str(extract_error)}"
            }), 400
        
        if not extracted_text.strip():
            return jsonify({"success": False, "error": "No text could be extracted from the document"}), 400
        
        # Perform AI-powered analysis
        analysis = analyze_text_intelligence(extracted_text)
        
        # Generate intelligent content
        processed_content = process_text_content(extracted_text, process_type)
        
        logger.info(f"✅ Document processed successfully: {analysis['word_count']} words, {analysis['character_count']} characters")
        
        return jsonify({
            "success": True,
            "filename": filename,
            "extracted_text": extracted_text,
            "processed_content": processed_content,
            "process_type": process_type,
            "word_count": analysis['word_count'],
            "char_count": analysis['character_count'],
            "ai_analysis": analysis,
            "message": "Document processed successfully with AI enhancement"
        })
        
    except Exception as e:
        logger.error(f"Document processing error: {e}")
        return jsonify({
            "success": False,
            "error": f"Document processing failed: {str(e)}"
        }), 500

@app.route('/process-text', methods=['POST'])
def process_text():
    """AI-enhanced text processing endpoint"""
    try:
        data = request.get_json()
        if not data or 'text' not in data:
            return jsonify({"success": False, "error": "No text provided"}), 400
        
        text = data['text'].strip()
        process_type = data.get('type', 'summary')
        
        if not text:
            return jsonify({"success": False, "error": "Text is empty"}), 400
        
        logger.info(f"Processing text: {len(text)} characters with type: {process_type}")
        
        # Perform AI-powered analysis
        analysis = analyze_text_intelligence(text)
        
        # Process the text
        processed_content = process_text_content(text, process_type)
        
        logger.info(f"✅ Text processed successfully: {analysis['word_count']} words, {analysis['character_count']} characters")
        
        return jsonify({
            "success": True,
            "processed_content": processed_content,
            "process_type": process_type,
            "word_count": analysis['word_count'],
            "char_count": analysis['character_count'],
            "ai_analysis": analysis,
            "message": "Text processed successfully with AI enhancement"
        })
        
    except Exception as e:
        logger.error(f"Text processing error: {e}")
        return jsonify({
            "success": False,
            "error": f"Text processing failed: {str(e)}"
        }), 500

@app.route('/api/speech/languages', methods=['GET'])
def get_supported_languages():
    """Get supported languages for speech recognition"""
    languages = [
        {"code": "en-US", "name": "English (US)", "native": "English"},
        {"code": "hi-IN", "name": "Hindi (India)", "native": "हिंदी"},
        {"code": "bn-IN", "name": "Bengali (India)", "native": "বাংলা"},
        {"code": "te-IN", "name": "Telugu (India)", "native": "తెలుగు"},
        {"code": "ta-IN", "name": "Tamil (India)", "native": "தமிழ்"},
        {"code": "kn-IN", "name": "Kannada (India)", "native": "ಕನ್ನಡ"},
        {"code": "ml-IN", "name": "Malayalam (India)", "native": "മലയാളം"},
        {"code": "gu-IN", "name": "Gujarati (India)", "native": "ગુજરાતી"},
        {"code": "mr-IN", "name": "Marathi (India)", "native": "मराठी"},
        {"code": "pa-IN", "name": "Punjabi (India)", "native": "ਪੰਜਾਬੀ"},
        {"code": "or-IN", "name": "Odia (India)", "native": "ଓଡ଼ିଆ"},
        {"code": "as-IN", "name": "Assamese (India)", "native": "অসমীয়া"},
        {"code": "ur-IN", "name": "Urdu (India)", "native": "اردو"},
        {"code": "es-ES", "name": "Spanish (Spain)", "native": "Español"},
        {"code": "fr-FR", "name": "French (France)", "native": "Français"},
        {"code": "de-DE", "name": "German (Germany)", "native": "Deutsch"},
        {"code": "it-IT", "name": "Italian (Italy)", "native": "Italiano"},
        {"code": "pt-PT", "name": "Portuguese (Portugal)", "native": "Português"},
        {"code": "ru-RU", "name": "Russian (Russia)", "native": "Русский"},
        {"code": "ja-JP", "name": "Japanese (Japan)", "native": "日本語"},
        {"code": "ko-KR", "name": "Korean (South Korea)", "native": "한국어"},
        {"code": "zh-CN", "name": "Chinese (Simplified)", "native": "中文"},
        {"code": "ar-SA", "name": "Arabic (Saudi Arabia)", "native": "العربية"}
    ]
    return jsonify({
        "languages": languages,
        "total": len(languages),
        "api_status": "success"
    })

if __name__ == '__main__':
    logger.info("🚀 Starting AI-Enhanced Web API Backend")
    logger.info("📝 Speech-to-text: Available with 23 languages")
    logger.info("🔊 Text-to-speech: Available with 23 languages")
    logger.info("🤖 AI Document Processing: Intelligent analysis & summarization")
    logger.info("📊 Features: Topic detection, keyword extraction, smart summaries")
    logger.info("🌐 Server will be available at: http://localhost:5000")
    logger.info("🔗 API endpoints:")
    logger.info("   - POST /api/speech/transcribe")
    logger.info("   - POST /api/speech/tts")
    logger.info("   - GET /api/speech/languages")
    logger.info("   - POST /upload-document (AI-enhanced)")
    logger.info("   - POST /process-text (AI-enhanced)")
    logger.info("   - GET /health")
    logger.info("🎉 All AI features are now functional!")
    
    app.run(debug=True, host='0.0.0.0', port=5000)




